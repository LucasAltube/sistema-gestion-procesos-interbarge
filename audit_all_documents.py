#!/usr/bin/env python3

import os
from docx import Document
from pathlib import Path

# Complete inventory for auditing
DOCUMENTS_TO_AUDIT = [
    # ADMINISTRACION (11 docs)
    {"code": "PROC-ADM-001", "area": "Administracion", "file": "PROC-ADM-001 Gestión de Weekly Report V.3.docx"},
    {"code": "PROC-ADM-002", "area": "Administracion", "file": "PROC-ADM-002 Gestión de Proyecciones V.3.docx"},
    {"code": "PROC-ADM-003", "area": "Administracion", "file": "PROC-ADM-003 Gestión de Pagos V.3.docx"},
    {"code": "PROC-ADM-004", "area": "Administracion", "file": "PROC-ADM-004 Gestión de Dashboard V.3.docx"},
    {"code": "PROC-ADM-005", "area": "Administracion", "file": "PROC-ADM-005 Proyección de Caja V.3.docx"},
    {"code": "PROC-ADM-006", "area": "Administracion", "file": "PROC-ADM-006 Gestión de Presupuestación V.3.docx"},
    {"code": "PROC-ADM-007", "area": "Administracion", "file": "PROC-ADM-007 Gestión de Cierre Contable V.3.docx"},
    {"code": "PROC-ADM-008", "area": "Administracion", "file": "PROC-ADM-008 Conciliación de Proveedores V.3.docx"},
    {"code": "PROC-ADM-009", "area": "Administracion", "file": "PROC-ADM-009 Gestión de Facturación V.3.docx"},
    {"code": "PROC-ADM-010", "area": "Administracion", "file": "PROC-ADM-010 Gestión de Cobranzas V.3.docx"},
    {"code": "INST-ADM-001", "area": "Administracion", "file": "INST-ADM-001 Instructivo de Facturación en sistema Netsuite V.3.docx"},
    
    # COMERCIAL (7 docs)
    {"code": "PROC-COM-001", "area": "Comercial", "file": "PROC-COM-001 Análisis de Mercado V.3.docx"},
    {"code": "PROC-COM-002", "area": "Comercial", "file": "PROC-COM-002 Gestión de Acuerdos Comerciales V.3.docx"},
    {"code": "PROC-COM-003", "area": "Comercial", "file": "PROC-COM-003 Comunicación interna y con Cliente V.3.docx"},
    {"code": "PROC-COM-004", "area": "Comercial", "file": "PROC-COM-004 Coordinación de operaciones portuarias V.3.docx"},
    {"code": "PROC-COM-005", "area": "Comercial", "file": "PROC-COM-005 Análisis para la Facturación Comercial V.3.docx"},
    {"code": "PROC-COM-006", "area": "Comercial", "file": "PROC-COM-006 Reportes Comerciales V.3.docx"},
    {"code": "INST-COM-001", "area": "Comercial", "file": "INST-COM-001 Instructivo Carga de Nuevo Proyecto V.3.docx"},
    
    # OPERACIONES (4 docs)
    {"code": "PROC-OPS-002", "area": "Operaciones", "file": "PROC-OPS-002 Programación de Viaje V.3.docx"},
    {"code": "PROC-OPS-003", "area": "Operaciones", "file": "PROC-OPS-003 Gestión de Operaciones V.3.docx"},
    {"code": "PROC-OPS-004", "area": "Operaciones", "file": "PROC-OPS-004 Gestión de Finalización de Viaje V.3.docx"},
    {"code": "PROC-OPS-005", "area": "Operaciones", "file": "PROC-OPS-005 Pronóstico y monitoreo de precipitaciones V.3.docx"},
    
    # RRHH (12 docs) - Updated to include PROC-RHU-003
    {"code": "PROC-RHU-003", "area": "RRHH", "file": "PROC-RHU-003 Reclutamiento selección y contratación del Personal V.3.docx"},
    {"code": "PROC-RHU-004", "area": "RRHH", "file": "PROC-RHU-004 Administrar el Bienestar de los Trabajadores V.3.docx"},
    {"code": "PROC-RHU-006", "area": "RRHH", "file": "PROC-RHU-006 Administración del Personal V.3.docx"},
    {"code": "PROC-RHU-007", "area": "RRHH", "file": "PROC-RHU-007 Administración de Seguro Médico V.3.docx"},
    {"code": "PROC-RHU-008", "area": "RRHH", "file": "PROC-RHU-008 Suspensión de Contratos de Trabajo V.3.docx"},
    {"code": "PROC-RHU-009", "area": "RRHH", "file": "PROC-RHU-009 Gestión de Convocatoria a Tripulación V.3.docx"},
    {"code": "PROC-RHU-011", "area": "RRHH", "file": "PROC-RHU-011 Información para liquidación Sueldos V.3.docx"},
    {"code": "PROC-RHU-012", "area": "RRHH", "file": "PROC-RHU-012 Gestión de insumos y mantenimiento V.3.docx"},
    {"code": "PROC-RHU-014", "area": "RRHH", "file": "PROC-RHU-014 Gestión Documental en oficina V.3.docx"},
    {"code": "PROC-RHU-015", "area": "RRHH", "file": "PROC-RHU-015 Gestión de viajes V.3.docx"},
    {"code": "PROC-RHU-016", "area": "RRHH", "file": "PROC-RHU-016 Evaluación de Desempeño V.3.docx"},
    {"code": "PROC-RHU-017", "area": "RRHH", "file": "PROC-RHU-017 Gestión de Uso de Vehículos V.3.docx"},
    
    # SCH (3 docs)
    {"code": "PROC-SCH-001", "area": "SCH", "file": "PROC-SCH-001 Gestión de Compras V.3.docx"},
    {"code": "PROC-SCH-002", "area": "SCH", "file": "PROC-SCH-002 Selección Evaluación y Control de Proveedores V.3.docx"},
    {"code": "PROC-SCH-003", "area": "SCH", "file": "PROC-SCH-003 Recepción de bienes y verificación de servicios V.3.docx"},
    
    # IT (1 doc)
    {"code": "PROC-IT-001", "area": "IT", "file": "PROC-IT-001 Mantenimiento de Infraestructura IT V.3.docx"},
    
    # SEGURIDAD (2 docs)
    {"code": "PROC-BUQ-003", "area": "Seguridad", "file": "PROC-BUQ-003 Plan de Emergencia V.3.docx"},
    {"code": "PROC-GES-001", "area": "Seguridad", "file": "PROC-GES-001 Gestión de Certificados V.3.docx"},
    
    # TECNICA (1 doc)
    {"code": "PROC-TEC-001", "area": "Tecnica", "file": "PROC-TEC-001 Mantenimiento de flota V.3.docx"},
]

def audit_document_issues(doc_path, doc_code):
    """Audit a single document for common issues"""
    issues = []
    
    try:
        doc = Document(doc_path)
        
        # 1. Check for excessive "Ver registro asociado"
        ver_registro_count = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "Ver registro asociado" in cell.text:
                        ver_registro_count += 1
        
        if ver_registro_count > 5:
            issues.append(f"❌ CRITICAL: {ver_registro_count} instances of 'Ver registro asociado' (likely spam)")
        elif ver_registro_count > 0:
            issues.append(f"⚠️  {ver_registro_count} instances of 'Ver registro asociado'")
        
        # 2. Check table sizes
        large_tables = []
        for i, table in enumerate(doc.tables):
            rows = len(table.rows)
            if rows > 20:
                large_tables.append(f"Table {i+1}: {rows} rows")
        
        if large_tables:
            issues.append(f"📊 Large tables detected: {', '.join(large_tables)}")
        
        # 3. Check for empty tables with just headers
        empty_tables_with_spam = 0
        for table in doc.tables:
            if len(table.rows) > 5:  # More than just a few header rows
                non_empty_content_rows = 0
                spam_rows = 0
                
                for row_idx in range(1, len(table.rows)):  # Skip header
                    row = table.rows[row_idx]
                    row_text = " ".join([cell.text for cell in row.cells])
                    if row_text.strip():
                        if "Ver registro asociado" in row_text:
                            spam_rows += 1
                        else:
                            non_empty_content_rows += 1
                
                if spam_rows > non_empty_content_rows and spam_rows > 5:
                    empty_tables_with_spam += 1
        
        if empty_tables_with_spam > 0:
            issues.append(f"💥 CRITICAL: {empty_tables_with_spam} tables with spam filling")
        
        # 4. Count total paragraphs and tables
        para_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        
        # 5. Look for font inconsistencies (basic check)
        font_changes = 0
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.name and run.font.name != "Calibri":  # Assuming Calibri is standard
                    font_changes += 1
                    break
        
        if font_changes > 0:
            issues.append(f"🔤 Font variations detected in {font_changes} paragraphs")
        
        return {
            'status': 'ERROR' if any('CRITICAL' in issue for issue in issues) else 'WARNING' if issues else 'OK',
            'issues': issues,
            'stats': {
                'paragraphs': para_count,
                'tables': table_count,
                'ver_registro_count': ver_registro_count
            }
        }
        
    except Exception as e:
        return {
            'status': 'FAIL',
            'issues': [f"💀 FAILED TO OPEN: {str(e)}"],
            'stats': {}
        }

def audit_all_documents():
    """Audit all processed documents"""
    base_dir = "/Users/lucas/Library/CloudStorage/GoogleDrive-altubelucas@gmail.com/Mi unidad/08 - ALTUBE IA/Interbarge/Sistema de Gestion por Procesos/Propuesta de Mejora Segura"
    
    print("🔍 STARTING COMPREHENSIVE AUDIT OF ALL PROCESSED DOCUMENTS")
    print("=" * 80)
    
    total_docs = len(DOCUMENTS_TO_AUDIT)
    ok_count = 0
    warning_count = 0
    error_count = 0
    fail_count = 0
    
    critical_issues = []
    
    for doc_info in DOCUMENTS_TO_AUDIT:
        doc_path = os.path.join(base_dir, doc_info['area'], doc_info['file'])
        
        print(f"\n📄 AUDITING: {doc_info['code']}")
        print(f"   Path: {doc_info['area']}/{doc_info['file']}")
        
        if not os.path.exists(doc_path):
            print("   ❌ FILE NOT FOUND")
            fail_count += 1
            critical_issues.append(f"{doc_info['code']}: FILE NOT FOUND")
            continue
        
        audit_result = audit_document_issues(doc_path, doc_info['code'])
        status = audit_result['status']
        
        if status == 'OK':
            print("   ✅ DOCUMENT OK")
            ok_count += 1
        elif status == 'WARNING':
            print("   ⚠️  DOCUMENT HAS WARNINGS")
            warning_count += 1
        elif status == 'ERROR':
            print("   ❌ DOCUMENT HAS CRITICAL ISSUES")
            error_count += 1
            critical_issues.append(f"{doc_info['code']}: CRITICAL ISSUES")
        else:  # FAIL
            print("   💀 DOCUMENT FAILED TO PROCESS")
            fail_count += 1
            critical_issues.append(f"{doc_info['code']}: PROCESSING FAILURE")
        
        # Print issues
        for issue in audit_result['issues']:
            print(f"     {issue}")
        
        # Print stats
        stats = audit_result['stats']
        if stats:
            print(f"     📊 Stats: {stats['paragraphs']} paragraphs, {stats['tables']} tables")
    
    print("\n" + "=" * 80)
    print("🎯 AUDIT SUMMARY")
    print("=" * 80)
    print(f"📊 Total documents audited: {total_docs}")
    print(f"✅ OK: {ok_count}")
    print(f"⚠️  Warnings: {warning_count}")
    print(f"❌ Critical Issues: {error_count}")
    print(f"💀 Failed to process: {fail_count}")
    print(f"📈 Success rate: {((ok_count + warning_count) / total_docs * 100):.1f}%")
    
    if critical_issues:
        print(f"\n🚨 CRITICAL ISSUES FOUND:")
        for issue in critical_issues:
            print(f"   {issue}")
    
    print(f"\n{'🎉 AUDIT COMPLETED SUCCESSFULLY' if error_count == 0 and fail_count == 0 else '⚠️ ISSUES DETECTED - CORRECTIVE ACTION NEEDED'}")
    
    return {
        'total': total_docs,
        'ok': ok_count,
        'warnings': warning_count,
        'errors': error_count,
        'failures': fail_count,
        'critical_issues': critical_issues
    }

if __name__ == "__main__":
    audit_all_documents()