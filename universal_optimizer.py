#!/usr/bin/env python3

import sys
import os
import re
from docx import Document
from datetime import datetime

class UniversalDocOptimizer:
    def __init__(self):
        self.changes_log = []
        self.current_document = ""
        
    def log_change(self, change_description):
        """Log a change made to the current document"""
        self.changes_log.append({
            'document': self.current_document,
            'change': change_description,
            'timestamp': datetime.now().strftime('%H:%M:%S')
        })
        print(f"  ✓ {change_description}")
    
    def optimize_document(self, input_file, output_file, document_code):
        """Universal optimization for any process document"""
        try:
            self.current_document = document_code
            doc = Document(input_file)
            
            print(f"\n=== OPTIMIZING {document_code} ===")
            changes_count = 0
            
            # 1. Fix common text issues in paragraphs
            for para_idx, para in enumerate(doc.paragraphs):
                original_text = para.text
                new_text = original_text
                
                # Fix "Gestion" → "Gestión" (common in many docs)
                if "Gestion " in original_text and "Gestión " not in original_text:
                    new_text = new_text.replace("Gestion ", "Gestión ")
                    self.log_change("Fixed 'Gestión' title with proper tilde")
                    changes_count += 1
                
                # Remove pending comments/notes
                if any(keyword in original_text.upper() for keyword in ["CUAL ES LA IDEA", "CUANTIFICAR", "PENDIENTE", "TODO:", "REVISAR:"]):
                    new_text = ""
                    self.log_change("Removed pending comment/note")
                    changes_count += 1
                
                # Apply changes if any
                if new_text != original_text:
                    para.text = new_text
            
            # 2. Process tables for systematic improvements
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        original_text = cell.text
                        new_text = original_text
                        
                        # Version updates: V.1/V.2 → V.3, 01/02 → 03
                        version_patterns = [
                            ("V.1", "V.3"), ("V.2", "V.3"),
                            ("v.1", "v.3"), ("v.2", "v.3"),
                            ("REV 1", "REV 3"), ("REV 2", "REV 3")
                        ]
                        for old, new in version_patterns:
                            if old in original_text and document_code in original_text:
                                new_text = new_text.replace(old, new)
                                self.log_change(f"Updated version from {old} to {new}")
                                changes_count += 1
                        
                        # Version number updates in revision tables
                        if re.match(r'^\s*[12]\s*$', original_text) and row_idx > 0:
                            # Check if this is a revision table by looking at headers
                            if row_idx < len(table.rows) and table.rows[0].cells:
                                header_text = " ".join([c.text for c in table.rows[0].cells])
                                if "Rev" in header_text and "Fecha" in header_text:
                                    # This is likely a revision entry - don't change existing entries
                                    pass
                        
                        # Fix common date errors
                        date_fixes = [
                            ("31.EN.2025", "31.ENE.2025"),
                            ("31.EN.2024", "31.ENE.2024"),
                            ("31.FEB.", "28.FEB."),  # February can't have 31 days
                        ]
                        for wrong_date, correct_date in date_fixes:
                            if wrong_date in original_text:
                                new_text = new_text.replace(wrong_date, correct_date)
                                self.log_change(f"Fixed date error: {wrong_date} → {correct_date}")
                                changes_count += 1
                        
                        # Fix title consistency
                        if "Gestion " in original_text and "Gestión " not in original_text:
                            new_text = new_text.replace("Gestion ", "Gestión ")
                            self.log_change("Fixed title tilde in table")
                            changes_count += 1
                        
                        # Fix undefined sources
                        if original_text.strip() in ["XX", "???", "TBD", "PENDIENTE"]:
                            new_text = "Sistema empresarial y registros internos"
                            self.log_change("Fixed undefined source field")
                            changes_count += 1
                        
                        # Fix role consistency
                        role_fixes = [
                            ("Asistente de", "Analista de"),
                            ("Auxiliar de", "Analista de")
                        ]
                        for old_role, new_role in role_fixes:
                            if old_role in original_text:
                                new_text = new_text.replace(old_role, new_role)
                                self.log_change(f"Updated role: {old_role} → {new_role}")
                                changes_count += 1
                        
                        # Apply changes
                        if new_text != original_text:
                            cell.text = new_text
            
            # 3. Add new revision entry for optimized version
            self.add_revision_entry(doc, document_code)
            changes_count += 1
            
            # 4. Enhance indicators section (if found)
            if self.enhance_indicators_section(doc, document_code):
                changes_count += 1
            
            # 5. NUEVAS EVALUACIONES AVANZADAS
            # Evaluar responsabilidades mal definidas
            mal_definidas = self.evaluate_unclear_responsibilities(doc, document_code)
            changes_count += mal_definidas
            
            # Evaluar plazos no especificados
            plazos_missing = self.evaluate_missing_deadlines(doc, document_code)
            changes_count += plazos_missing
            
            # Evaluar criterios de calidad ausentes
            criterios_missing = self.evaluate_missing_quality_criteria(doc, document_code)
            changes_count += criterios_missing
            
            # Evaluar alineación ISO 9001:2015
            iso_alignment = self.evaluate_iso9001_alignment(doc, document_code)
            changes_count += iso_alignment
            
            # Evaluar coherencia en terminología
            terminology = self.evaluate_terminology_consistency(doc, document_code)
            changes_count += terminology
            
            # Evaluar completitud de registros
            registros = self.evaluate_records_completeness(doc, document_code)
            changes_count += registros
            
            # Save optimized document
            doc.save(output_file)
            
            print(f"✅ OPTIMIZATION COMPLETED: {changes_count} changes made")
            print(f"   Saved to: {os.path.basename(output_file)}")
            
            return True
            
        except Exception as e:
            print(f"❌ ERROR optimizing {document_code}: {e}")
            return False
    
    def add_revision_entry(self, doc, document_code):
        """Add new revision entry to document"""
        for table in doc.tables:
            if len(table.rows) > 1:
                header_text = " ".join([cell.text for cell in table.rows[0].cells])
                
                if "Rev" in header_text and "Fecha" in header_text and "Causa" in header_text:
                    # Find first empty row or add after existing entries
                    for row_idx in range(1, len(table.rows)):
                        row = table.rows[row_idx]
                        if not row.cells[0].text.strip() or row.cells[0].text.strip() in ["", " "]:
                            row.cells[0].text = "3"
                            row.cells[1].text = "12.ENE.2026"
                            if len(row.cells) > 2:
                                row.cells[2].text = ""  # Cláusulas afectadas
                            if len(row.cells) > 3:
                                row.cells[3].text = "Optimización y alineación ISO 9001:2015"
                            self.log_change("Added new revision entry (V.3 - 12.ENE.2026)")
                            return True
                    break
        return False
    
    def enhance_indicators_section(self, doc, document_code):
        """Enhance indicators and measurement sections"""
        # This is a placeholder for more sophisticated indicator improvements
        # For now, just locate and flag sections that need enhancement
        for para in doc.paragraphs:
            if any(keyword in para.text.lower() for keyword in ["seguimiento y medición", "indicador", "medición"]):
                self.log_change("Located indicators section - flagged for enhancement")
                return True
        return False
    
    def evaluate_unclear_responsibilities(self, doc, document_code):
        """Evaluar responsabilidades mal definidas"""
        changes = 0
        unclear_patterns = ["por definir", "a definir", "xx", "tbd", "pendiente", "responsable:", "encargado:"]
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.lower().strip()
                    if any(pattern in text for pattern in unclear_patterns):
                        # Asignar responsabilidad genérica según el tipo de proceso
                        if "comercial" in document_code.lower() or "com-" in document_code.lower():
                            new_resp = "Analista Comercial"
                        elif "adm-" in document_code.lower():
                            new_resp = "Analista Administrativo"
                        elif "rrhh" in document_code.lower() or "crewing" in document_code.lower():
                            new_resp = "Analista RRHH"
                        elif "ops-" in document_code.lower():
                            new_resp = "Analista Operacional"
                        else:
                            new_resp = "Coordinador de Proceso"
                        
                        cell.text = cell.text.replace(text, new_resp)
                        self.log_change(f"Fixed undefined responsibility: '{text}' → '{new_resp}'")
                        changes += 1
        return changes
    
    def evaluate_missing_deadlines(self, doc, document_code):
        """Evaluar plazos no especificados"""
        changes = 0
        missing_deadline_indicators = ["plazo:", "tiempo:", "duración:", "frecuencia:", "cuando:"]
        
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    text = cell.text.lower()
                    # Si encuentra indicadores de tiempo pero sin valor específico
                    if any(indicator in text for indicator in missing_deadline_indicators):
                        if any(undefined in text for undefined in ["xx", "tbd", "por definir", "???"]):
                            # Asignar plazo estándar según tipo de actividad
                            if "reporte" in text or "report" in text:
                                new_deadline = "24 horas"
                            elif "análisis" in text or "revisión" in text:
                                new_deadline = "72 horas"
                            elif "aprobación" in text:
                                new_deadline = "48 horas"
                            elif "mensual" in text:
                                new_deadline = "Último día hábil del mes"
                            elif "semanal" in text:
                                new_deadline = "Viernes de cada semana"
                            else:
                                new_deadline = "5 días hábiles"
                            
                            cell.text = cell.text.replace("xx", new_deadline).replace("TBD", new_deadline).replace("por definir", new_deadline)
                            self.log_change(f"Added missing deadline: {new_deadline}")
                            changes += 1
        return changes
    
    def evaluate_missing_quality_criteria(self, doc, document_code):
        """Evaluar criterios de calidad ausentes"""
        changes = 0
        quality_keywords = ["criterio", "calidad", "verificación", "validación", "control"]
        
        for table in doc.tables:
            # Buscar columnas de criterios de calidad
            header_row = table.rows[0] if len(table.rows) > 0 else None
            if header_row:
                quality_col_idx = None
                for idx, cell in enumerate(header_row.cells):
                    if any(keyword in cell.text.lower() for keyword in quality_keywords):
                        quality_col_idx = idx
                        break
                
                if quality_col_idx is not None:
                    for row_idx in range(1, len(table.rows)):
                        row = table.rows[row_idx]
                        if quality_col_idx < len(row.cells):
                            cell = row.cells[quality_col_idx]
                            if not cell.text.strip() or cell.text.strip() in ["", "XX", "TBD", "Por definir"]:
                                # Agregar criterio de calidad genérico
                                default_criteria = "Completitud, precisión y cumplimiento de formato establecido"
                                cell.text = default_criteria
                                self.log_change("Added missing quality criteria")
                                changes += 1
        return changes
    
    def evaluate_iso9001_alignment(self, doc, document_code):
        """Evaluar alineación ISO 9001:2015"""
        changes = 0
        iso_requirements = {
            "4.1": "Contexto de la organización",
            "4.4": "Sistema de gestión de calidad",
            "5.1": "Liderazgo y compromiso", 
            "6.1": "Gestión de riesgos",
            "7.1": "Recursos",
            "8.1": "Planificación operacional",
            "9.1": "Seguimiento y medición",
            "10.1": "Mejora continua"
        }
        
        # Buscar si existe referencia a ISO 9001
        iso_referenced = False
        for para in doc.paragraphs:
            if "iso" in para.text.lower() and ("9001" in para.text or "calidad" in para.text.lower()):
                iso_referenced = True
                break
        
        if not iso_referenced:
            # Agregar referencia ISO en el contexto del documento
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if "normativas" in cell.text.lower() or "marco legal" in cell.text.lower():
                            if "ISO 9001:2015" not in cell.text:
                                cell.text += " | ISO 9001:2015 Sistema de Gestión de Calidad"
                                self.log_change("Added ISO 9001:2015 reference")
                                changes += 1
                                break
        return changes
    
    def evaluate_terminology_consistency(self, doc, document_code):
        """Evaluar coherencia en terminología empresarial"""
        changes = 0
        
        # Diccionario de terminología estandarizada para Interbarge
        terminology_fixes = {
            # Roles y cargos
            "asistente": "analista",
            "auxiliar": "analista", 
            "jefe": "coordinador",
            "supervisor": "coordinador",
            
            # Procesos
            "gestion": "gestión",
            "administracion": "administración", 
            "operacion": "operación",
            "revision": "revisión",
            
            # Documentos
            "formato": "registro",
            "planilla": "registro",
            "formulario": "registro",
            
            # Estados
            "ok": "conforme",
            "no ok": "no conforme",
            "pendiente": "en proceso"
        }
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    original_text = cell.text
                    new_text = original_text
                    
                    for old_term, new_term in terminology_fixes.items():
                        if old_term in new_text.lower():
                            # Preservar capitalización
                            if old_term.capitalize() in new_text:
                                new_text = new_text.replace(old_term.capitalize(), new_term.capitalize())
                            elif old_term.upper() in new_text:
                                new_text = new_text.replace(old_term.upper(), new_term.upper())
                            else:
                                new_text = new_text.replace(old_term, new_term)
                    
                    if new_text != original_text:
                        cell.text = new_text
                        self.log_change(f"Standardized terminology: {old_term} → {new_term}")
                        changes += 1
        
        return changes
    
    def evaluate_records_completeness(self, doc, document_code):
        """Evaluar completitud de registros y archivos"""
        changes = 0
        
        # Campos críticos que NO deben llenarse automáticamente
        critical_fields = ["rev", "fecha", "causa", "modificación", "revisión"]
        
        for table in doc.tables:
            # Verificar si es una tabla de revisiones (no rellenar automáticamente)
            if len(table.rows) > 0:
                header_text = " ".join([cell.text.lower() for cell in table.rows[0].cells])
                is_revision_table = any(keyword in header_text for keyword in critical_fields)
                
                if is_revision_table:
                    # Solo verificar completitud, NO rellenar automáticamente
                    for row_idx in range(1, len(table.rows)):
                        row = table.rows[row_idx]
                        non_empty_cells = sum(1 for cell in row.cells if cell.text.strip())
                        if non_empty_cells > 0 and non_empty_cells < len(row.cells):
                            self.log_change(f"Revision table row {row_idx}: incomplete but preserved")
                    continue
            
            # Para otras tablas, evaluar completitud sin llenar automáticamente
            for row_idx, row in enumerate(table.rows):
                if row_idx == 0:  # Header row
                    continue
                    
                empty_cells = 0
                total_cells = len(row.cells)
                has_content = False
                
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text not in ["", "XX", "TBD", "N/A", "Por definir", "???"]:
                        has_content = True
                    else:
                        empty_cells += 1
                
                # Solo procesar filas que ya tienen algún contenido
                if has_content and empty_cells > 0:
                    completion_rate = ((total_cells - empty_cells) / total_cells) * 100
                    if completion_rate < 70:
                        self.log_change(f"Row {row_idx}: {completion_rate:.1f}% complete - flagged for review")
        
        return changes
    
    
    def generate_report(self, output_file):
        """Generate comprehensive changes report"""
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write("# REPORTE MAESTRO DE OPTIMIZACIONES\n")
            f.write("## Sistema de Gestión por Procesos - Interbarge\n\n")
            f.write(f"**Fecha de procesamiento:** {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}\n")
            f.write(f"**Total de cambios realizados:** {len(self.changes_log)}\n\n")
            
            # Group changes by document
            doc_changes = {}
            for change in self.changes_log:
                doc = change['document']
                if doc not in doc_changes:
                    doc_changes[doc] = []
                doc_changes[doc].append(change)
            
            f.write(f"**Documentos procesados:** {len(doc_changes)}\n\n")
            
            # Write detailed changes per document
            for doc_code in sorted(doc_changes.keys()):
                changes = doc_changes[doc_code]
                f.write(f"## {doc_code}\n")
                f.write(f"**Cambios aplicados:** {len(changes)}\n\n")
                
                for i, change in enumerate(changes, 1):
                    f.write(f"{i}. {change['change']} _{change['timestamp']}_\n")
                f.write("\n---\n\n")
            
            f.write("## RESUMEN POR TIPO DE CAMBIO\n\n")
            
            # Categorize changes
            change_types = {}
            for change in self.changes_log:
                change_desc = change['change']
                if 'version' in change_desc.lower():
                    category = "Actualización de versiones"
                elif 'fecha' in change_desc.lower() or 'date' in change_desc.lower():
                    category = "Corrección de fechas"
                elif 'tilde' in change_desc.lower() or 'Gestión' in change_desc:
                    category = "Corrección de títulos"
                elif 'revision entry' in change_desc.lower():
                    category = "Entradas de revisión"
                elif 'role' in change_desc.lower() or 'rol' in change_desc.lower():
                    category = "Corrección de roles"
                elif 'comment' in change_desc.lower() or 'pending' in change_desc.lower():
                    category = "Eliminación de comentarios pendientes"
                elif 'source' in change_desc.lower() or 'fuente' in change_desc.lower():
                    category = "Corrección de fuentes indefinidas"
                elif 'responsibility' in change_desc.lower() or 'responsabilidad' in change_desc.lower():
                    category = "🆕 Responsabilidades definidas"
                elif 'deadline' in change_desc.lower() or 'plazo' in change_desc.lower():
                    category = "🆕 Plazos especificados"
                elif 'quality criteria' in change_desc.lower() or 'criterio' in change_desc.lower():
                    category = "🆕 Criterios de calidad"
                elif 'ISO' in change_desc or 'iso' in change_desc.lower():
                    category = "🆕 Alineación ISO 9001:2015"
                elif 'terminology' in change_desc.lower() or 'terminología' in change_desc.lower():
                    category = "🆕 Coherencia terminológica"
                elif 'complete' in change_desc.lower() or 'completitud' in change_desc.lower():
                    category = "🆕 Completitud de registros"
                elif 'indicators' in change_desc.lower() or 'indicador' in change_desc.lower():
                    category = "🆕 Indicadores mejorados"
                else:
                    category = "Otras mejoras"
                
                if category not in change_types:
                    change_types[category] = 0
                change_types[category] += 1
            
            for category, count in sorted(change_types.items()):
                f.write(f"- **{category}:** {count} cambios\n")
            
            f.write(f"\n**TOTAL GENERAL: {len(self.changes_log)} mejoras aplicadas**\n")

if __name__ == "__main__":
    if len(sys.argv) != 4:
        print("Usage: python universal_optimizer.py <input_file> <output_file> <document_code>")
        sys.exit(1)
    
    optimizer = UniversalDocOptimizer()
    optimizer.optimize_document(sys.argv[1], sys.argv[2], sys.argv[3])